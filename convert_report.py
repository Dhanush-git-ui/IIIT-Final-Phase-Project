import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Add standard Calibri font styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def add_inline_formatting(paragraph, text):
    # Parses basic bold **text** and adds runs
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Parse italic *text* or _text_
            sub_parts = re.split(r'(\*.*?\*|_.*?_)', part)
            for sub_part in sub_parts:
                if (sub_part.startswith('*') and sub_part.endswith('*')) or (sub_part.startswith('_') and sub_part.endswith('_')):
                    run = paragraph.add_run(sub_part[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(sub_part)

md_path = r"c:\Users\dhanu\OneDrive\Desktop\IIIT GRAND FINAL\Project\Project_Report.md"
docx_path = r"c:\Users\dhanu\OneDrive\Desktop\IIIT GRAND FINAL\Project\Project_Report.docx"

with open(md_path, 'r', encoding='utf-8') as f:
    lines = f.readlines()

table_lines = []
in_table = False

for line in lines:
    line_str = line.strip()
    
    # Check for tables
    if line_str.startswith('|'):
        in_table = True
        table_lines.append(line_str)
        continue
    elif in_table:
        # We finished the table, process it
        in_table = False
        if len(table_lines) > 1:
            # Parse table lines
            rows_data = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.split('|')[1:-1]]
                rows_data.append(cells)
            
            # Remove separator row like | --- | --- |
            if len(rows_data) > 1 and any(c.startswith('-') for c in rows_data[1]):
                rows_data.pop(1)
                
            if rows_data:
                num_rows = len(rows_data)
                num_cols = len(rows_data[0])
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.style = 'Light Shading Accent 1' # A standard word style
                for r_idx, row in enumerate(rows_data):
                    for c_idx, cell_val in enumerate(row):
                        cell = table.cell(r_idx, c_idx)
                        p = cell.paragraphs[0]
                        add_inline_formatting(p, cell_val)
                doc.add_paragraph() # spacing
        table_lines = []
    
    # Ignore empty lines
    if not line_str:
        continue
        
    # Headers
    if line_str.startswith('# '):
        doc.add_heading(line_str[2:], level=0)
    elif line_str.startswith('## '):
        doc.add_heading(line_str[3:], level=1)
    elif line_str.startswith('### '):
        doc.add_heading(line_str[4:], level=2)
    elif line_str.startswith('#### '):
        doc.add_heading(line_str[5:], level=3)
    # Horizontal line
    elif line_str == '---':
        p = doc.add_paragraph()
        p_border = p.add_run('__________________________________________________')
        p_border.font.color.rgb = RGBColor(148, 163, 184)
    # Bullet lists
    elif line_str.startswith('- ') or line_str.startswith('* '):
        bullet_text = line_str[2:]
        p = doc.add_paragraph(style='List Bullet')
        add_inline_formatting(p, bullet_text)
    # Numbered list item
    elif re.match(r'^\d+\.\s', line_str):
        match = re.match(r'^(\d+\.\s)(.*)', line_str)
        p = doc.add_paragraph(style='List Number')
        add_inline_formatting(p, match.group(2))
    # Normal paragraph
    else:
        p = doc.add_paragraph()
        add_inline_formatting(p, line_str)

# Double check if any table is left at the end of file
if in_table and len(table_lines) > 1:
    rows_data = []
    for tl in table_lines:
        cells = [c.strip() for c in tl.split('|')[1:-1]]
        rows_data.append(cells)
    if len(rows_data) > 1 and any(c.startswith('-') for c in rows_data[1]):
        rows_data.pop(1)
    if rows_data:
        num_rows = len(rows_data)
        num_cols = len(rows_data[0])
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Light Shading Accent 1'
        for r_idx, row in enumerate(rows_data):
            for c_idx, cell_val in enumerate(row):
                cell = table.cell(r_idx, c_idx)
                p = cell.paragraphs[0]
                add_inline_formatting(p, cell_val)

doc.save(docx_path)
print("Saved Docx Successfully to:", docx_path)
